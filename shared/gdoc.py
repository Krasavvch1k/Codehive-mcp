"""
Спільна логіка для роботи з .docx файлами та експортованими Google Docs.

Використовується у:
- projects.worqen.parsers.team_discussions (обговорення з командою)
- projects.codehive.gdoc_reader (документи у CodeHive Agency Drive)

Workflow: Google Doc експортується через Drive API у .docx (байти),
далі ці байти обробляються функціями цього модуля.
"""

import io
import re
from typing import Optional

from docx import Document


def docx_bytes_to_markdown(content: bytes) -> str:
    """
    Перетворює байти .docx у markdown-рендер.

    - Heading 1..6 → # ## ### #### ##### ######
    - Параграфи зі стилями → plain text
    - Таблиці → markdown-таблиці з роздільником "|"
    - Порожні параграфи пропускаються
    """
    docx_doc = Document(io.BytesIO(content))
    pieces: list[str] = []

    for para in docx_doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        m = re.match(r"^Heading\s+(\d+)$", style)
        if m:
            level = int(m.group(1))
            pieces.append(f"{'#' * level} {text}")
        else:
            pieces.append(text)

    for table in docx_doc.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            if len(rows) > 1:
                sep = " | ".join(["---"] * len(table.rows[0].cells))
                rows.insert(1, sep)
            pieces.append("\n".join(rows))

    return "\n\n".join(pieces)


def docx_bytes_to_plain_text(content: bytes) -> str:
    """
    Перетворює байти .docx у плоский текст (один параграф на рядок).

    Без форматування. Призначено для пошуку / індексації.
    """
    docx_doc = Document(io.BytesIO(content))
    return "\n".join(p.text for p in docx_doc.paragraphs)


def extract_snippet(text: str, query: str, context_chars: int = 200) -> Optional[str]:
    """
    Шукає query у text (case-insensitive). Якщо знайдено — повертає snippet
    з ±context_chars/2 символами контексту навколо знахідки. Snippet
    обрамлюється "..." на тих краях де відрізали від оригінального тексту.

    Якщо не знайдено — повертає None.
    """
    q_lower = query.lower()
    pos = text.lower().find(q_lower)
    if pos == -1:
        return None

    start = max(0, pos - context_chars // 2)
    end = min(len(text), pos + len(query) + context_chars // 2)
    snippet = text[start:end]
    if start > 0:
        snippet = "..." + snippet
    if end < len(text):
        snippet = snippet + "..."
    return snippet
