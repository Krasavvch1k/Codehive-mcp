"""Generic write handlers для нативних .docx файлів (НЕ Google Docs).

Низькорівневі функції — приймають уже resolved file_id і виконують операцію через
python-docx + redownload/modify/reupload через shared.drive_client. НЕ роблять:
- resolve query → file_id (це project-specific)
- safety check (blacklist/whitelist policy теж project-specific)

Проєктні обгортки повинні зробити resolve + safety ДО виклику, передати готовий
file_id/file_name.

Logging — спільний через shared.writes_log.log_doc_write, з параметром project.

Особливості docx vs gdoc:
- Нема revisionId-локінгу — використовуємо modifiedTime через check_drive_unchanged
- python-docx розбиває текст на runs (фрагменти з форматуванням). Простий алгоритм
  заміни: якщо old_text цілком у одному run — точна заміна зі збереженням стилю.
  Якщо через кілька runs — заміна з можливою втратою внутрішнього форматування у
  межах new_text (стилі границь runs зберігаються).
"""

import io
import logging
from typing import Any, Optional

from docx import Document
from googleapiclient.errors import HttpError

from shared.drive_client import (
    download_file,
    fetch_current_drive_modified,
    upload_file_content,
)
from shared.safety import SafetyError, check_drive_unchanged
from shared.writes_log import log_doc_write

logger = logging.getLogger(__name__)


# Mime для нативних .docx
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

CONTEXT_RADIUS = 30  # символів зліва/справа для error-контексту
VALID_INSERT_MODES = ("after", "before", "end_of_doc")


# ----- helpers -----

def _paragraphs_full_text(doc: "Document") -> tuple[list[str], list[Any]]:
    """Повертає (тексти параграфів, об'єкти параграфів) у порядку обходу.

    Обходить body + таблиці (їх клітинки теж містять параграфи).
    """
    texts: list[str] = []
    paragraphs: list[Any] = []

    for p in doc.paragraphs:
        texts.append(p.text)
        paragraphs.append(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    texts.append(p.text)
                    paragraphs.append(p)

    return texts, paragraphs


def _find_paragraph_occurrences(
    paragraphs_texts: list[str], needle: str
) -> list[tuple[int, int]]:
    """Повертає список (paragraph_index, char_offset_in_paragraph) усіх входжень.

    case-sensitive, всередині параграфа, без перетину між параграфами.
    """
    occurrences: list[tuple[int, int]] = []
    for p_idx, text in enumerate(paragraphs_texts):
        start = 0
        while True:
            i = text.find(needle, start)
            if i == -1:
                break
            occurrences.append((p_idx, i))
            start = i + 1
    return occurrences


def _format_paragraph_contexts(
    paragraphs_texts: list[str],
    occurrences: list[tuple[int, int]],
    needle: str,
) -> list[str]:
    """Контекст ±CONTEXT_RADIUS символів навколо кожного occurrence."""
    contexts: list[str] = []
    n = len(needle)
    for p_idx, offset in occurrences:
        text = paragraphs_texts[p_idx]
        left = text[max(0, offset - CONTEXT_RADIUS) : offset]
        right = text[offset + n : offset + n + CONTEXT_RADIUS]
        left_clean = " ".join(left.split())
        right_clean = " ".join(right.split())
        contexts.append(f"[p{p_idx}] ...{left_clean} [{needle}] {right_clean}...")
    return contexts


def _replace_in_paragraph(paragraph: Any, old_text: str, new_text: str) -> bool:
    """
    Замінює old_text на new_text у одному параграфі.

    Алгоритм:
    1. Якщо old_text цілком у одному run — replace всередині run, стиль run
       зберігається.
    2. Якщо old_text через кілька runs — конкатенуємо всі runs у paragraph,
       робимо string replace, очищаємо runs крім першого, перший run отримує
       весь новий текст (стиль перших runs зберігається, внутрішнього
       форматування у межах new_text не буде).

    Повертає True якщо заміна відбулась, False якщо old_text не знайдено.
    """
    full = paragraph.text
    if old_text not in full:
        return False

    # Спроба точної заміни всередині одного run
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text, 1)
            return True

    # old_text через кілька runs — мердж усього параграфа у перший run
    new_full = full.replace(old_text, new_text, 1)
    runs = paragraph.runs
    if runs:
        runs[0].text = new_full
        for run in runs[1:]:
            run.text = ""
    else:
        # Param без runs (edge case) — додаємо новий run
        paragraph.add_run(new_full)
    return True


def _insert_text_in_paragraph(
    paragraph: Any, anchor: str, text_to_insert: str, mode: str
) -> bool:
    """
    Вставка тексту у параграф relative до anchor.

    mode='after': text_to_insert йде після anchor
    mode='before': text_to_insert йде перед anchor

    Реалізація — через string-rebuild параграфа (як для multi-run replace).
    Стиль перших runs зберігається, у вставленому тексті форматування не буде.

    Повертає True якщо вставка відбулась, False якщо anchor не знайдено.
    """
    full = paragraph.text
    if anchor not in full:
        return False

    anchor_pos = full.find(anchor)
    if mode == "after":
        insert_pos = anchor_pos + len(anchor)
    elif mode == "before":
        insert_pos = anchor_pos
    else:
        raise ValueError(f"Unsupported mode for paragraph insert: {mode}")

    new_full = full[:insert_pos] + text_to_insert + full[insert_pos:]

    runs = paragraph.runs
    if runs:
        runs[0].text = new_full
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_full)
    return True


# ----- public API -----

def replace_text_in_docx(
    file_id: str,
    file_name: str,
    old_text: str,
    new_text: str,
    project: str,
    force_overwrite: bool = False,
) -> dict[str, Any]:
    """
    Замінити рівно ОДНЕ входження old_text на new_text у нативному .docx.

    Caller відповідає за resolve query → file_id і safety check ДО виклику.

    Args:
        file_id: Drive ID документа.
        file_name: назва (для логу і error messages).
        old_text: текст для пошуку (case-sensitive).
        new_text: чим замінити.
        project: "codehive" / "worqen" — для логу.
        force_overwrite: пропустити drive-unchanged check.

    Returns:
        dict з ok / error.
    """
    if not old_text:
        return {"error": "old_text не може бути порожнім"}

    if old_text == new_text:
        return {"error": "old_text == new_text, нічого замінювати"}

    # Optimistic locking — fetch modifiedTime ДО завантаження
    baseline_modified = fetch_current_drive_modified(file_id)

    # Read
    try:
        raw = download_file(file_id, fmt="docx", force_refresh=True)
        doc = Document(io.BytesIO(raw))
    except HttpError as e:
        return {"error": f"Не вдалось завантажити docx: {e}"}
    except Exception as e:
        return {"error": f"Не вдалось розпарсити docx: {e}"}

    # Count occurrences
    paragraphs_texts, paragraphs = _paragraphs_full_text(doc)
    occurrences = _find_paragraph_occurrences(paragraphs_texts, old_text)
    count = len(occurrences)

    if count == 0:
        return {
            "error": (
                f"Текст '{old_text}' не знайдено у '{file_name}'. "
                f"Перевір регістр (case-sensitive) і пробіли. "
                f"Зверни увагу: пошук іде всередині параграфів, не через переноси."
            ),
            "file_id": file_id,
            "file_name": file_name,
        }

    if count > 1:
        contexts = _format_paragraph_contexts(paragraphs_texts, occurrences, old_text)
        return {
            "error": (
                f"Знайдено {count} входжень '{old_text}' у '{file_name}'. "
                f"replace_text заміняє тільки якщо рівно 1 входження. "
                f"Зроби old_text більш специфічним (додай контекст з боків)."
            ),
            "file_id": file_id,
            "file_name": file_name,
            "occurrences": count,
            "contexts": contexts[:10],
        }

    # Drive sync check перед write
    if not force_overwrite:
        try:
            check_drive_unchanged(file_id, baseline_modified)
        except SafetyError as e:
            return {
                "error": str(e),
                "kind": e.kind,
                "file_id": file_id,
                "file_name": file_name,
                "hint": "Передай force_overwrite=True щоб ігнорувати.",
            }

    # Write — у потрібному параграфі
    p_idx, _offset = occurrences[0]
    target_paragraph = paragraphs[p_idx]
    replaced = _replace_in_paragraph(target_paragraph, old_text, new_text)
    if not replaced:
        # Не повинно статись після перевірки occurrence — параноя
        return {
            "error": (
                "Заміна не відбулась всередині параграфа попри знайдене входження. "
                "Можливо текст розбитий нестандартно. Спробуй коротший old_text."
            ),
            "file_id": file_id,
            "file_name": file_name,
        }

    # Serialize + upload
    try:
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        upload_meta = upload_file_content(file_id, out.getvalue(), DOCX_MIME)
    except HttpError as e:
        return {
            "error": f"Drive upload failed: {e}",
            "file_id": file_id,
            "file_name": file_name,
        }

    drive_modified_after = upload_meta.get("modifiedTime")

    # Log
    try:
        log_doc_write(
            project=project,
            tool="replace_text_docx",
            file_id=file_id,
            file_name=file_name,
            payload={
                "old_text": old_text,
                "new_text": new_text,
                "paragraph_index": p_idx,
                "modified_before": baseline_modified,
                "modified_after": drive_modified_after,
                "force_overwrite": force_overwrite,
            },
        )
    except Exception as e:
        logger.warning("writes_log failed for %s.replace_text_docx: %s", project, e)

    return {
        "ok": True,
        "file_id": file_id,
        "file_name": file_name,
        "paragraph_index": p_idx,
        "modified_before": baseline_modified,
        "modified_after": drive_modified_after,
        "replaced": 1,
        "force_overwrite_used": force_overwrite,
    }


def insert_text_in_docx(
    file_id: str,
    file_name: str,
    text: str,
    mode: str,
    anchor: Optional[str],
    project: str,
    force_overwrite: bool = False,
) -> dict[str, Any]:
    """
    Вставка тексту у нативний .docx у режимі after/before/end_of_doc.

    Args:
        file_id: Drive ID документа.
        file_name: назва.
        text: що вставляти.
        mode: "after" | "before" | "end_of_doc".
        anchor: для after/before — substring (case-sensitive, unique у одному
            параграфі і у документі). Для end_of_doc — має бути None або порожній.
        project: "codehive" / "worqen" — для логу.
        force_overwrite: пропустити drive-unchanged check.

    Для after/before — anchor має бути в межах одного параграфа.
    Для end_of_doc — text додається новим параграфом у кінець body.

    Returns:
        dict з ok / error.
    """
    # Validate
    if mode not in VALID_INSERT_MODES:
        return {
            "error": f"mode має бути одним з {VALID_INSERT_MODES}, отримано: '{mode}'"
        }

    if not text:
        return {"error": "text не може бути порожнім"}

    if mode in ("after", "before"):
        if not anchor:
            return {"error": f"для mode='{mode}' потрібен anchor"}
    else:  # end_of_doc
        if anchor:
            return {
                "error": (
                    "для mode='end_of_doc' anchor має бути None або порожнім; "
                    "якщо треба вставити після конкретного тексту — використовуй mode='after'"
                )
            }

    baseline_modified = fetch_current_drive_modified(file_id)

    # Read
    try:
        raw = download_file(file_id, fmt="docx", force_refresh=True)
        doc = Document(io.BytesIO(raw))
    except HttpError as e:
        return {"error": f"Не вдалось завантажити docx: {e}"}
    except Exception as e:
        return {"error": f"Не вдалось розпарсити docx: {e}"}

    inserted_paragraph_idx: Optional[int] = None

    if mode == "end_of_doc":
        # Додаємо новий параграф у кінець body (не у таблиці).
        # python-docx: doc.add_paragraph() додає у самий кінець body.
        new_para = doc.add_paragraph(text)
        # Індекс — len(doc.paragraphs) - 1 у новому стані
        inserted_paragraph_idx = len(doc.paragraphs) - 1
    else:
        # after / before — anchor у тексті
        paragraphs_texts, paragraphs = _paragraphs_full_text(doc)
        occurrences = _find_paragraph_occurrences(paragraphs_texts, anchor)
        count = len(occurrences)

        if count == 0:
            return {
                "error": (
                    f"Anchor '{anchor}' не знайдено у '{file_name}'. "
                    f"Перевір регістр (case-sensitive) і пробіли."
                ),
                "file_id": file_id,
                "file_name": file_name,
            }

        if count > 1:
            contexts = _format_paragraph_contexts(paragraphs_texts, occurrences, anchor)
            return {
                "error": (
                    f"Знайдено {count} входжень '{anchor}' у '{file_name}'. "
                    f"insert_text вимагає рівно 1 входження. "
                    f"Зроби anchor більш специфічним."
                ),
                "file_id": file_id,
                "file_name": file_name,
                "occurrences": count,
                "contexts": contexts[:10],
            }

        p_idx, _offset = occurrences[0]
        target_paragraph = paragraphs[p_idx]
        inserted = _insert_text_in_paragraph(target_paragraph, anchor, text, mode)
        if not inserted:
            return {
                "error": (
                    "Вставка не відбулась попри знайдений anchor. "
                    "Можливо anchor розбитий нестандартно. Спробуй коротший anchor."
                ),
                "file_id": file_id,
                "file_name": file_name,
            }
        inserted_paragraph_idx = p_idx

    # Drive sync check перед write
    if not force_overwrite:
        try:
            check_drive_unchanged(file_id, baseline_modified)
        except SafetyError as e:
            return {
                "error": str(e),
                "kind": e.kind,
                "file_id": file_id,
                "file_name": file_name,
                "hint": "Передай force_overwrite=True щоб ігнорувати.",
            }

    # Serialize + upload
    try:
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        upload_meta = upload_file_content(file_id, out.getvalue(), DOCX_MIME)
    except HttpError as e:
        return {
            "error": f"Drive upload failed: {e}",
            "file_id": file_id,
            "file_name": file_name,
        }

    drive_modified_after = upload_meta.get("modifiedTime")

    # Log
    try:
        log_doc_write(
            project=project,
            tool="insert_text_docx",
            file_id=file_id,
            file_name=file_name,
            payload={
                "mode": mode,
                "anchor": anchor,
                "text": text,
                "inserted_paragraph_idx": inserted_paragraph_idx,
                "modified_before": baseline_modified,
                "modified_after": drive_modified_after,
                "force_overwrite": force_overwrite,
            },
        )
    except Exception as e:
        logger.warning("writes_log failed for %s.insert_text_docx: %s", project, e)

    return {
        "ok": True,
        "file_id": file_id,
        "file_name": file_name,
        "mode": mode,
        "inserted_paragraph_idx": inserted_paragraph_idx,
        "modified_before": baseline_modified,
        "modified_after": drive_modified_after,
        "force_overwrite_used": force_overwrite,
    }


def create_docx_in_folder(
    folder_id: str,
    folder_name: str,
    name: str,
    initial_content: str,
    project: str,
) -> dict[str, Any]:
    """
    Створює новий .docx у вказаній папці.

    Caller відповідає за resolve folder_query → folder_id, duplicate check
    і safety check ДО виклику.

    Args:
        folder_id: Drive ID папки куди створити.
        folder_name: назва папки (для логу і response).
        name: назва нового документа (без .docx — додамо самі якщо потрібно).
        initial_content: опційно, plain text для тіла (кожен \\n — новий параграф).
        project: "codehive" / "worqen" — для логу.

    Returns:
        dict з ok / error.
    """
    from shared.drive_client import get_service

    # Створюємо docx у пам'яті
    doc = Document()
    if initial_content:
        for line in initial_content.split("\n"):
            doc.add_paragraph(line)
    out = io.BytesIO()
    doc.save(out)
    content_bytes = out.getvalue()

    # Додаємо .docx якщо нема
    if not name.lower().endswith(".docx"):
        upload_name = name + ".docx"
    else:
        upload_name = name

    service = get_service()

    from googleapiclient.http import MediaIoBaseUpload

    media = MediaIoBaseUpload(
        io.BytesIO(content_bytes),
        mimetype=DOCX_MIME,
        resumable=False,
    )

    try:
        new_file = (
            service.files()
            .create(
                body={
                    "name": upload_name,
                    "mimeType": DOCX_MIME,
                    "parents": [folder_id],
                },
                media_body=media,
                fields="id, name, mimeType, parents, webViewLink, modifiedTime",
                supportsAllDrives=True,
            )
            .execute()
        )
    except HttpError as e:
        return {"error": f"Drive create failed: {e}"}

    file_id = new_file["id"]
    web_view_link = new_file.get("webViewLink") or (
        f"https://drive.google.com/file/d/{file_id}/view"
    )

    # Log
    try:
        log_doc_write(
            project=project,
            tool="create_docx",
            file_id=file_id,
            file_name=upload_name,
            payload={
                "folder_id": folder_id,
                "folder_name": folder_name,
                "initial_content_length": len(initial_content),
                "modified_after": new_file.get("modifiedTime"),
            },
        )
    except Exception as e:
        logger.warning("writes_log failed for %s.create_docx: %s", project, e)

    return {
        "ok": True,
        "file_id": file_id,
        "file_name": upload_name,
        "folder_id": folder_id,
        "folder_name": folder_name,
        "url": web_view_link,
        "modified_after": new_file.get("modifiedTime"),
        "initial_content_length": len(initial_content),
    }
