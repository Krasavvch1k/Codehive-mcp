"""Read operations for the CodeHive Agency Drive folder."""

import io
import re
from typing import Optional

from docx import Document

from drive_client import (
    _list_folder_children,
    download_file,
    _get_service,
)
from projects.codehive.config import (
    CODEHIVE_ROOT_FOLDER_ID,
    GOOGLE_DOC_MIME,
    GOOGLE_FOLDER_MIME,
    CODEHIVE_MAX_RECURSION_DEPTH,
)


def _resolve_root(folder_id: Optional[str]) -> str:
    if folder_id:
        return folder_id
    if not CODEHIVE_ROOT_FOLDER_ID:
        raise RuntimeError(
            "CODEHIVE_ROOT_FOLDER_ID is not set in .env. "
            "Add CODEHIVE_ROOT_FOLDER_ID=<id> and restart the MCP server."
        )
    return CODEHIVE_ROOT_FOLDER_ID


def _classify_item(item: dict) -> str:
    mime = item.get("mimeType", "")
    if mime == GOOGLE_FOLDER_MIME:
        return "folder"
    if mime == GOOGLE_DOC_MIME:
        return "gdoc"
    if mime == "application/vnd.google-apps.spreadsheet":
        return "gsheet"
    if mime == "application/vnd.google-apps.presentation":
        return "gslides"
    if mime == "application/pdf":
        return "pdf"
    if mime.startswith("image/"):
        return "image"
    return "other"


def _normalize_item(item: dict, parent_id: str, parent_name: Optional[str]) -> dict:
    return {
        "id": item["id"],
        "name": item.get("name", ""),
        "mime": item.get("mimeType", ""),
        "kind": _classify_item(item),
        "modified": item.get("modifiedTime"),
        "size_bytes": int(item["size"]) if item.get("size") else None,
        "parent_folder_id": parent_id,
        "parent_folder_name": parent_name,
    }


def list_folder(folder_id: Optional[str] = None) -> dict:
    fid = _resolve_root(folder_id)

    service = _get_service()
    try:
        folder_meta = service.files().get(
            fileId=fid,
            fields="id, name, mimeType",
            supportsAllDrives=True,
        ).execute()
        folder_name = folder_meta.get("name")
    except Exception:
        folder_name = None

    raw_children = _list_folder_children(fid)

    folders: list[dict] = []
    docs: list[dict] = []
    for item in raw_children:
        normalized = _normalize_item(item, fid, folder_name)
        if normalized["kind"] == "folder":
            folders.append(normalized)
        else:
            docs.append(normalized)

    folders.sort(key=lambda d: d["name"].lower())
    docs.sort(key=lambda d: d["name"].lower())

    return {
        "folder_id": fid,
        "folder_name": folder_name,
        "folders_count": len(folders),
        "docs_count": len(docs),
        "folders": folders,
        "docs": docs,
    }


def list_all_docs(max_depth: int = CODEHIVE_MAX_RECURSION_DEPTH) -> dict:
    root_id = _resolve_root(None)
    service = _get_service()
    try:
        root_meta = service.files().get(
            fileId=root_id,
            fields="id, name",
            supportsAllDrives=True,
        ).execute()
        root_name = root_meta.get("name")
    except Exception:
        root_name = None

    items: list[dict] = []
    total_folders = 0
    total_docs = 0

    queue: list[tuple[str, Optional[str], str, int]] = [
        (root_id, root_name, "", 0),
    ]

    while queue:
        cur_id, cur_name, cur_path, depth = queue.pop(0)
        if depth > max_depth:
            continue

        children = _list_folder_children(cur_id)
        for item in children:
            normalized = _normalize_item(item, cur_id, cur_name)
            sub_path = (
                f"{cur_path} > {normalized['name']}"
                if cur_path
                else normalized["name"]
            )
            normalized["path"] = sub_path
            normalized["depth"] = depth + 1
            items.append(normalized)

            if normalized["kind"] == "folder":
                total_folders += 1
                if depth + 1 < max_depth:
                    queue.append((normalized["id"], normalized["name"], sub_path, depth + 1))
            else:
                total_docs += 1

    items.sort(key=lambda d: d["path"].lower())

    return {
        "root_folder_id": root_id,
        "root_folder_name": root_name,
        "max_depth": max_depth,
        "total_folders": total_folders,
        "total_docs": total_docs,
        "items": items,
    }


def _resolve_doc(query: str, all_docs: list[dict]) -> dict:
    q = query.strip()

    if len(q) > 25 and " " not in q and "(" not in q:
        for d in all_docs:
            if d["id"] == q:
                return d

    q_lower = q.lower()
    matches = [d for d in all_docs if q_lower in d["name"].lower()]

    if not matches:
        raise ValueError(
            f"Document matching '{query}' not found in CodeHive Agency. "
            f"Try codehive_list_all_docs to see all names."
        )

    if len(matches) > 1:
        names = [
            f"{m['name']} [{m.get('path', '?')}] (id: {m['id'][:12]}...)"
            for m in matches[:5]
        ]
        more = f", ...{len(matches) - 5} more" if len(matches) > 5 else ""
        raise ValueError(
            f"Query '{query}' matched {len(matches)} documents. "
            f"Refine the query. Candidates: {'; '.join(names)}{more}"
        )

    return matches[0]


def read_doc(query: str) -> dict:
    all_data = list_all_docs(max_depth=CODEHIVE_MAX_RECURSION_DEPTH)
    gdocs = [d for d in all_data["items"] if d["kind"] == "gdoc"]

    doc_meta = _resolve_doc(query, gdocs)

    content = download_file(doc_meta["id"], fmt="gdoc")
    docx_doc = Document(io.BytesIO(content))

    pieces: list[str] = []
    for para in docx_doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        m = re.match(r"^Heading\s+(\d+)$", style)
        if m:
            level = int(m.group(1))
            pieces.append(f"{'#' * level} {text}")
        else:
            pieces.append(text)

    for table in docx_doc.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            if len(rows) > 1:
                sep = " | ".join(["---"] * len(table.rows[0].cells))
                rows.insert(1, sep)
            pieces.append("\n".join(rows))

    text = "\n\n".join(pieces)

    return {
        "id": doc_meta["id"],
        "name": doc_meta["name"],
        "mime": doc_meta["mime"],
        "path": doc_meta.get("path"),
        "modified": doc_meta.get("modified"),
        "length": len(text),
        "text": text,
    }


def search(
    query: str,
    scope: str = "names",
    limit: int = 20,
    context_chars: int = 200,
) -> dict:
    if scope not in ("names", "content", "both"):
        raise ValueError(f"scope must be 'names' / 'content' / 'both', not '{scope}'")

    all_data = list_all_docs(max_depth=CODEHIVE_MAX_RECURSION_DEPTH)
    items = all_data["items"]
    q_lower = query.lower()

    name_matches: list[dict] = []
    content_matches: list[dict] = []

    if scope in ("names", "both"):
        for it in items:
            if q_lower in it["name"].lower() or (
                it.get("path") and q_lower in it["path"].lower()
            ):
                name_matches.append(it)

    if scope in ("content", "both"):
        gdocs = [d for d in items if d["kind"] == "gdoc"]
        for d in gdocs:
            try:
                content = download_file(d["id"], fmt="gdoc")
                docx_doc = Document(io.BytesIO(content))
                full_text = "\n".join(p.text for p in docx_doc.paragraphs)
                pos = full_text.lower().find(q_lower)
                if pos != -1:
                    start = max(0, pos - context_chars // 2)
                    end = min(len(full_text), pos + len(query) + context_chars // 2)
                    snippet = full_text[start:end]
                    if start > 0:
                        snippet = "..." + snippet
                    if end < len(full_text):
                        snippet = snippet + "..."
                    content_matches.append({
                        "id": d["id"],
                        "name": d["name"],
                        "path": d.get("path"),
                        "snippet": snippet,
                    })
            except Exception as e:
                content_matches.append({
                    "id": d["id"],
                    "name": d["name"],
                    "path": d.get("path"),
                    "error": str(e),
                })

    return {
        "scope": scope,
        "query": query,
        "total_name_matches": len(name_matches),
        "total_content_matches": len(content_matches),
        "name_matches": name_matches[:limit],
        "content_matches": content_matches[:limit],
    }
