"""
Універсальний парсер для doc-документів Worqen
(Tech Doc, PRD Bootstrap, PRD v1.1, PRD v2, AML Policy, ToS, Privacy, Cookie).

Розбирає документ на дерево секцій по Heading-стилях.
Підтримує: повний текст, конкретну секцію, TOC з breadcrumbs,
блок-між-двома-заголовками, пошук, таблиці.
"""

import io
import re
from typing import Optional
from docx import Document
from docx.document import Document as _DocumentType
from docx.table import Table
from docx.oxml.ns import qn

from drive_client import download_file
from projects.worqen.config import FILE_IDS, FILE_FORMATS


# Ключі FILE_IDS які є doc-документами (включно з gdoc — він експортується у docx).
DOC_KEYS = [
    "tech_doc",
    "prd_bootstrap",
    "prd_v1_1",
    "prd_v2",
    "aml_policy",
    "tos",
    "privacy",
    "cookie",
]


def _heading_level(style_name: Optional[str]) -> Optional[int]:
    """Повертає рівень заголовку 1-9 або None якщо це не Heading."""
    if not style_name:
        return None
    m = re.match(r"^Heading\s+(\d+)$", style_name)
    if m:
        return int(m.group(1))
    return None


def _table_to_text(table: Table) -> str:
    """Конвертує таблицю docx у markdown-подібний текст."""
    lines = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        lines.append(" | ".join(cells))
    if not lines:
        return ""
    if len(lines) > 1:
        sep = " | ".join(["---"] * len(table.rows[0].cells))
        lines.insert(1, sep)
    return "\n".join(lines)


def _iter_block_items(doc: _DocumentType):
    """
    Ітерує параграфи і таблиці у порядку появи у документі.
    """
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag
        if tag == qn("w:p"):
            for para in doc.paragraphs:
                if para._element is child:
                    yield ("p", para)
                    break
        elif tag == qn("w:tbl"):
            for table in doc.tables:
                if table._element is child:
                    yield ("t", table)
                    break


def _load_doc(file_key: str) -> _DocumentType:
    """
    Завантажує документ за ключем. Підтримує і нативні docx, і Google Docs
    (для gdoc drive_client експортує у docx через export_media).
    """
    if file_key not in FILE_IDS:
        raise ValueError(f"Невідомий ключ: {file_key}. Доступні: {DOC_KEYS}")
    if file_key not in DOC_KEYS:
        raise ValueError(f"{file_key} не є doc-документом. Доступні: {DOC_KEYS}")

    fmt = FILE_FORMATS.get(file_key, "docx")
    content = download_file(FILE_IDS[file_key], fmt=fmt)
    return Document(io.BytesIO(content))


def list_doc_keys() -> list[str]:
    """Список доступних doc-документів."""
    return DOC_KEYS


def _build_breadcrumb(stack: list[tuple[int, str]], current_title: str) -> str:
    """
    Будує breadcrumb-шлях типу '3. Architecture > 3.2 Backend > 3.2.1 FastAPI'
    зі стеку батьківських заголовків.
    """
    parts = [title for _lvl, title in stack]
    parts.append(current_title)
    return " > ".join(parts)


def get_toc(file_key: str) -> list[dict]:
    """
    Повертає Table of Contents — список усіх заголовків з рівнями і breadcrumbs.
    Кожен елемент:
        {
            "index": int,
            "level": int,    # 1-9
            "title": str,
            "path":  str,    # breadcrumb від кореня до цього заголовку
        }
    """
    doc = _load_doc(file_key)
    toc: list[dict] = []
    idx = 0

    # Стек відкритих батьківських заголовків: [(level, title), ...]
    stack: list[tuple[int, str]] = []

    for para in doc.paragraphs:
        level = _heading_level(para.style.name if para.style else None)
        if not level:
            continue
        title = para.text.strip()
        if not title:
            continue

        # Закриваємо у стеку всі заголовки з level >= поточного
        while stack and stack[-1][0] >= level:
            stack.pop()

        path = _build_breadcrumb(stack, title)

        toc.append({
            "index": idx,
            "level": level,
            "title": title,
            "path": path,
        })

        stack.append((level, title))
        idx += 1

    return toc


def get_section(
    file_key: str, heading_query: str, include_subsections: bool = True
) -> Optional[dict]:
    """
    Повертає секцію документу — від заголовка що матчить heading_query
    до наступного заголовка такого ж або вищого рівня.
    """
    doc = _load_doc(file_key)
    body_children = list(doc.element.body.iterchildren())

    start_idx = None
    start_level = None
    start_title = None

    for i, child in enumerate(body_children):
        if child.tag != qn("w:p"):
            continue
        for para in doc.paragraphs:
            if para._element is child:
                level = _heading_level(para.style.name if para.style else None)
                title = para.text.strip()
                if level and heading_query.lower() in title.lower():
                    start_idx = i
                    start_level = level
                    start_title = title
                    break
        if start_idx is not None:
            break

    if start_idx is None:
        return None

    pieces: list[str] = []

    for i in range(start_idx + 1, len(body_children)):
        child = body_children[i]

        if child.tag == qn("w:p"):
            for para in doc.paragraphs:
                if para._element is child:
                    level = _heading_level(para.style.name if para.style else None)
                    if level:
                        if level <= start_level:
                            return {
                                "title": start_title,
                                "level": start_level,
                                "text": "\n\n".join(pieces).strip(),
                            }
                        if not include_subsections:
                            return {
                                "title": start_title,
                                "level": start_level,
                                "text": "\n\n".join(pieces).strip(),
                            }
                        pieces.append(f"{'#' * level} {para.text.strip()}")
                    else:
                        text = para.text.strip()
                        if text:
                            pieces.append(text)
                    break

        elif child.tag == qn("w:tbl"):
            for table in doc.tables:
                if table._element is child:
                    table_text = _table_to_text(table)
                    if table_text:
                        pieces.append(table_text)
                    break

    return {
        "title": start_title,
        "level": start_level,
        "text": "\n\n".join(pieces).strip(),
    }


def _find_heading_body_idx(
    body_children: list, doc: _DocumentType, query: str
) -> Optional[tuple[int, int, str]]:
    """
    Знаходить перший заголовок який матчить query (case-insensitive substring).
    Повертає (body_index, level, title) або None.
    """
    query_lower = query.lower()
    for i, child in enumerate(body_children):
        if child.tag != qn("w:p"):
            continue
        for para in doc.paragraphs:
            if para._element is child:
                level = _heading_level(para.style.name if para.style else None)
                title = para.text.strip()
                if level and query_lower in title.lower():
                    return (i, level, title)
                break
    return None


def get_doc_block(
    file_key: str,
    heading_from: str,
    heading_to: Optional[str] = None,
) -> Optional[dict]:
    """
    Повертає блок документа від heading_from (включно) до heading_to (включно).
    Якщо heading_to не вказано — від heading_from до кінця документа.

    На відміну від get_section, не зважає на рівні заголовків — бере все
    що між двома вказаними якорями.

    Параметри:
        file_key:     ключ документа
        heading_from: повний текст або частина заголовку-початку (case-insensitive).
                      Беремо ПЕРШИЙ матч.
        heading_to:   повний текст або частина заголовку-кінця (case-insensitive),
                      опціонально. Якщо вказано — береться ПЕРШИЙ матч ПІСЛЯ
                      heading_from. Сам цей заголовок включається до результату;
                      його контент обмежується наступним заголовком будь-якого
                      рівня.

    Повертає:
        {
            "from_title": str,
            "to_title":   str | None,
            "text":       str,
        }
        або None якщо heading_from не знайдено.
    """
    doc = _load_doc(file_key)
    body_children = list(doc.element.body.iterchildren())

    found_from = _find_heading_body_idx(body_children, doc, heading_from)
    if found_from is None:
        return None
    from_idx, _from_level, from_title = found_from

    to_idx: Optional[int] = None
    to_title: Optional[str] = None

    if heading_to is not None:
        to_query_lower = heading_to.lower()
        for j in range(from_idx + 1, len(body_children)):
            child = body_children[j]
            if child.tag != qn("w:p"):
                continue
            for para in doc.paragraphs:
                if para._element is child:
                    level = _heading_level(para.style.name if para.style else None)
                    title = para.text.strip()
                    if level and to_query_lower in title.lower():
                        to_idx = j
                        to_title = title
                    break
            if to_idx is not None:
                break

    pieces: list[str] = []
    end_idx = len(body_children)
    capture_tail_after_to = to_idx is not None

    for i in range(from_idx, end_idx):
        child = body_children[i]

        if child.tag == qn("w:p"):
            for para in doc.paragraphs:
                if para._element is child:
                    level = _heading_level(para.style.name if para.style else None)
                    text = para.text.strip()

                    # Якщо вже пройшли to_idx — перериваємо на наступному заголовку
                    if (
                        capture_tail_after_to
                        and i > to_idx
                        and level
                    ):
                        return {
                            "from_title": from_title,
                            "to_title": to_title,
                            "text": "\n\n".join(pieces).strip(),
                        }

                    if level and text:
                        pieces.append(f"{'#' * level} {text}")
                    elif text:
                        pieces.append(text)
                    break

        elif child.tag == qn("w:tbl"):
            for table in doc.tables:
                if table._element is child:
                    table_text = _table_to_text(table)
                    if table_text:
                        pieces.append(table_text)
                    break

    return {
        "from_title": from_title,
        "to_title": to_title,
        "text": "\n\n".join(pieces).strip(),
    }


def get_full_text(file_key: str) -> str:
    """
    Повертає весь текст документу як markdown.
    """
    doc = _load_doc(file_key)
    pieces: list[str] = []

    for kind, item in _iter_block_items(doc):
        if kind == "p":
            level = _heading_level(item.style.name if item.style else None)
            text = item.text.strip()
            if not text:
                continue
            if level:
                pieces.append(f"{'#' * level} {text}")
            else:
                pieces.append(text)
        elif kind == "t":
            table_text = _table_to_text(item)
            if table_text:
                pieces.append(table_text)

    return "\n\n".join(pieces)


def search_in_doc(
    file_key: str, query: str, context_chars: int = 200
) -> list[dict]:
    """
    Пошук підрядка у документі. Повертає список матчів з контекстом.
    """
    doc = _load_doc(file_key)
    query_lower = query.lower()
    matches = []
    current_section = "(beginning)"

    for kind, item in _iter_block_items(doc):
        if kind == "p":
            level = _heading_level(item.style.name if item.style else None)
            text = item.text.strip()
            if level and text:
                current_section = text
                continue
            if not text:
                continue

            text_lower = text.lower()
            pos = text_lower.find(query_lower)
            if pos != -1:
                start = max(0, pos - context_chars // 2)
                end = min(len(text), pos + len(query) + context_chars // 2)
                snippet = text[start:end]
                if start > 0:
                    snippet = "..." + snippet
                if end < len(text):
                    snippet = snippet + "..."
                matches.append({"section": current_section, "snippet": snippet})

        elif kind == "t":
            for row in item.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if query_lower in cell_text.lower():
                        snippet = cell_text
                        if len(snippet) > context_chars:
                            snippet = snippet[:context_chars] + "..."
                        matches.append(
                            {
                                "section": current_section + " (table)",
                                "snippet": snippet,
                            }
                        )

    return matches
