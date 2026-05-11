"""
Діагностичний скрипт: показує структуру docx файлу.
Виводить усі заголовки (heading-стилі) і їх рівні.
"""

import sys
import io
from docx import Document

from shared.drive_client import download_file
from projects.worqen.config import FILE_IDS


def inspect(file_key: str):
    if file_key not in FILE_IDS:
        print(f"Невідомий ключ: {file_key}")
        print(f"Доступні ключі: {list(FILE_IDS.keys())}")
        return

    print(f"Інспектую файл: {file_key}")
    print("=" * 60)

    file_id = FILE_IDS[file_key]
    content = download_file(file_id)
    doc = Document(io.BytesIO(content))

    total_paragraphs = len(doc.paragraphs)
    total_tables = len(doc.tables)
    print(f"Параграфів: {total_paragraphs}")
    print(f"Таблиць: {total_tables}")
    print()

    print("Усі заголовки (heading-стилі):")
    print("-" * 60)
    headings_count = 0
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            text = para.text.strip()
            if text:
                indent = "  " * (int(style_name.replace("Heading ", "")) - 1)
                print(f"{indent}[{style_name}] {text}")
                headings_count += 1

    print()
    print(f"Всього заголовків: {headings_count}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Використання: python inspect_docx.py <file_key>")
        print(f"Доступні: {[k for k, v in FILE_IDS.items() if v]}")
        sys.exit(1)

    inspect(sys.argv[1])
